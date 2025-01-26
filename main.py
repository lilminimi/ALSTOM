from flask import Flask, request, render_template
import os
import docx
import fitz  # PyMuPDF
import pandas as pd
import re
from werkzeug.utils import secure_filename
from typing import Dict, Any, Optional

app = Flask(__name__)

app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
app.config['ALLOWED_EXTENSIONS'] = {'docx', 'pdf', 'xlsx', 'xls', 'xltm'}


class AlstomDocumentVerifier:
    def __init__(self, template_path: Optional[str] = None):
        self.template_path = template_path
        self.confidentiality_levels = ['Public', 'Restricted', 'Confidential', 'Secret']
        self.document_numbering_patterns = {
            'central': r'^[A-Z]{3}-[A-Z]{3}-\d{3}$',
            'country': r'^[A-Z]{3}-[A-Z]{2}-[A-Z]{3}-\d{3}$',
            'site_product_line': r'^[A-Z]{3}-[A-Z]{2}-[A-Z]{3}-[A-Z]{3}-\d{3}$'
        }

    def verify_document(self, document_path: str, language: str) -> Dict[str, Any]:
        file_extension = document_path.split('.')[-1].lower()

        if file_extension == 'docx':
            return self._verify_word(document_path, language)
        elif file_extension == 'pdf':
            return self._verify_pdf(document_path, language)
        elif file_extension in {'xlsx', 'xls', 'xltm'}:
            return self._verify_excel(document_path, language)
        else:
            return {"error": "Unsupported file format."}

    def _verify_word(self, document_path: str, language: str) -> Dict[str, Any]:
        if not os.path.exists(document_path):
            return {"error": f"File not found at '{document_path}'"}

        try:
            doc = docx.Document(document_path)
            full_text = '\n'.join([p.text for p in doc.paragraphs])

            return {
                'title_present': self._format_result(self._check_title(doc), "Document must have a meaningful title.",
                                                     "Dokument muss einen sinnvollen Titel haben.",
                                                     doc.core_properties.title, language),
                'identification_number': self._format_result(self._check_identification_number(full_text),
                                                             "Document must follow identification numbering patterns.",
                                                             "Dokument muss Identifikationsnummerierungsrichtlinien folgen.",
                                                             re.findall(r'^[A-Z]{3}-[A-Z]{3}-\d{3}', full_text),
                                                             language),
                'page_numbers_correct': self._format_result(self._check_page_numbers(doc),
                                                            "Each page should include page numbers.",
                                                            "Jede Seite sollte Seitenzahlen enthalten.",
                                                            "Page numbers detected" if self._check_page_numbers(
                                                                doc) else "No page numbers found", language),
                'revision_status': self._format_result(self._check_revision_status(full_text),
                                                       "Document must indicate a revision status.",
                                                       "Dokument muss einen Versionsstatus angeben.",
                                                       re.findall(r'Version:?[\sA-D]', full_text, re.IGNORECASE),
                                                       language),
                'author_verified': self._format_result(self._check_author(doc), "Document must specify an author.",
                                                       "Dokument muss einen Autor angeben.", doc.core_properties.author,
                                                       language),
                'approval_verified': self._format_result(self._check_approval(full_text),
                                                         "Approval markers must be present.",
                                                         "Genehmigungsmarker müssen vorhanden sein.",
                                                         re.findall(r'approved|validated|freigegeben', full_text,
                                                                    re.IGNORECASE), language),
                'effective_date': self._format_result(self._check_effective_date(full_text),
                                                      "Document must include an effective date.",
                                                      "Dokument muss ein Gültigkeitsdatum enthalten.",
                                                      re.findall(r'\d{2}\.\d{2}\.\d{4}|\d{4}-\d{2}-\d{2}', full_text),
                                                      language),
                'ownership_notice': self._format_result(self._check_ownership_notice(full_text),
                                                        "Ownership notices must be included.",
                                                        "Eigentumshinweise müssen enthalten sein.",
                                                        re.findall(r'© ALSTOM|confidential|geschäftsgeheimnis',
                                                                   full_text, re.IGNORECASE), language),
                'creating_unit': self._format_result(self._check_creating_unit(full_text),
                                                     "Creating unit must be mentioned.",
                                                     "Erstellende Einheit muss erwähnt werden.",
                                                     re.findall(r'alstom|produktlinie|standort|dach region', full_text,
                                                                re.IGNORECASE), language),
                'confidentiality_level': self._format_result(self._check_confidentiality_level(full_text) is not None,
                                                             "Confidentiality level must be indicated.",
                                                             "Vertraulichkeitsstufe muss angegeben werden.",
                                                             self._check_confidentiality_level(full_text), language),
                'document_numbering': self._format_result(self._check_document_numbering(full_text),
                                                          "Document must follow numbering conventions.",
                                                          "Dokument muss Nummerierungskonventionen folgen.",
                                                          re.findall(r'[A-Z]{3}-[A-Z]{3}-\d{3}', full_text), language),
                'template_compliance': self._format_result(self._check_template_compliance(doc),
                                                           "Document must comply with the template.",
                                                           "Dokument muss dem Template entsprechen.",
                                                           "Template used" if self._check_template_compliance(
                                                               doc) else "No template detected", language),
                'ams_compliance': self._format_result(self._check_ams_compliance(full_text),
                                                      "Document must comply with AMS.",
                                                      "Dokument muss AMS entsprechen.",
                                                      re.findall(r'AMS|management handbook', full_text, re.IGNORECASE),
                                                      language)
            }
        except Exception as e:
            return {"error": f"Error during Word document verification: {e}"}

    def _verify_pdf(self, document_path: str, language: str) -> Dict[str, Any]:
        try:
            doc = fitz.open(document_path)
            full_text = '\n'.join(page.get_text() for page in doc)

            return {
                'page_count': self._format_result(doc.page_count > 0, "PDF must have pages.", "PDF muss Seiten haben.",
                                                  doc.page_count, language),
                'identification_number': self._format_result(self._check_identification_number(full_text),
                                                             "Document must follow identification numbering patterns.",
                                                             "Dokument muss Identifikationsnummerierungsrichtlinien folgen.",
                                                             re.findall(r'^[A-Z]{3}-[A-Z]{3}-\d{3}', full_text),
                                                             language),
                'revision_status': self._format_result(self._check_revision_status(full_text),
                                                       "Document must indicate a revision status.",
                                                       "Dokument muss einen Versionsstatus angeben.",
                                                       re.findall(r'Version:?[\sA-D]', full_text, re.IGNORECASE),
                                                       language),
                'approval_verified': self._format_result(self._check_approval(full_text),
                                                         "Approval markers must be present.",
                                                         "Genehmigungsmarker müssen vorhanden sein.",
                                                         re.findall(r'approved|validated|freigegeben', full_text,
                                                                    re.IGNORECASE), language),
                'effective_date': self._format_result(self._check_effective_date(full_text),
                                                      "Document must include an effective date.",
                                                      "Dokument muss ein Gültigkeitsdatum enthalten.",
                                                      re.findall(r'\d{2}\.\d{2}\.\d{4}|\d{4}-\d{2}-\d{2}', full_text),
                                                      language),
                'ownership_notice': self._format_result(self._check_ownership_notice(full_text),
                                                        "Ownership notices must be included.",
                                                        "Eigentumshinweise müssen enthalten sein.",
                                                        re.findall(r'© ALSTOM|confidential|geschäftsgeheimnis',
                                                                   full_text, re.IGNORECASE), language),
                'creating_unit': self._format_result(self._check_creating_unit(full_text),
                                                     "Creating unit must be mentioned.",
                                                     "Erstellende Einheit muss erwähnt werden.",
                                                     re.findall(r'alstom|produktlinie|standort|dach region', full_text,
                                                                re.IGNORECASE), language),
                'confidentiality_level': self._format_result(self._check_confidentiality_level(full_text) is not None,
                                                             "Confidentiality level must be indicated.",
                                                             "Vertraulichkeitsstufe muss angegeben werden.",
                                                             self._check_confidentiality_level(full_text), language),
                'document_numbering': self._format_result(self._check_document_numbering(full_text),
                                                          "Document must follow numbering conventions.",
                                                          "Dokument muss Nummerierungskonventionen folgen.",
                                                          re.findall(r'[A-Z]{3}-[A-Z]{3}-\d{3}', full_text), language),
                'ams_compliance': self._format_result(self._check_ams_compliance(full_text),
                                                      "Document must comply with AMS.",
                                                      "Dokument muss AMS entsprechen.",
                                                      re.findall(r'AMS|management handbook', full_text, re.IGNORECASE),
                                                      language)
            }
        except Exception as e:
            return {"error": f"Error during PDF verification: {e}"}

    def _verify_excel(self, document_path: str, language: str) -> Dict[str, Any]:
        try:
            df = pd.read_excel(document_path, engine='openpyxl')
            full_text = ' '.join(df.astype(str).values.flatten())

            return {
                'rows_present': self._format_result(not df.empty, "Excel file must have rows.",
                                                    "Excel-Datei muss Zeilen enthalten.", f"Rows: {len(df)}", language),
                'columns_present': self._format_result(len(df.columns) > 0, "Excel file must have columns.",
                                                       "Excel-Datei muss Spalten enthalten.",
                                                       f"Columns: {len(df.columns)}", language),
                'confidentiality_level': self._format_result(self._check_confidentiality_level(full_text) is not None,
                                                             "Confidentiality level must be indicated.",
                                                             "Vertraulichkeitsstufe muss angegeben werden.",
                                                             self._check_confidentiality_level(full_text), language),
                'ownership_notice': self._format_result(self._check_ownership_notice(full_text),
                                                        "Ownership notices must be included.",
                                                        "Eigentumshinweise müssen enthalten sein.",
                                                        re.findall(r'© ALSTOM|confidential|geschäftsgeheimnis',
                                                                   full_text, re.IGNORECASE), language),
                'creating_unit': self._format_result(self._check_creating_unit(full_text),
                                                     "Creating unit must be mentioned.",
                                                     "Erstellende Einheit muss erwähnt werden.",
                                                     re.findall(r'alstom|produktlinie|standort|dach region', full_text,
                                                                re.IGNORECASE), language)
            }
        except Exception as e:
            return {"error": f"Error during Excel verification: {e}"}

    def _format_result(self, passed, detail_en, detail_de, content, language):
        details = detail_en if language == 'en' else detail_de
        return {'passed': passed, 'details': details, 'content': content}

    def _check_title(self, doc: docx.Document) -> bool:
        title = doc.core_properties.title
        return bool(title and len(title) > 5)

    def _check_identification_number(self, text: str) -> bool:
        patterns = list(self.document_numbering_patterns.values())
        return any(re.search(pattern, text) for pattern in patterns)

    def _check_page_numbers(self, doc: docx.Document) -> bool:
        for para in doc.paragraphs:
            if re.search(r'Page \d+ of \d+', para.text) or re.search(r'\d+/\d+', para.text):
                return True
        return False

    def _check_revision_status(self, text: str) -> bool:
        return bool(re.search(r'Version:?[\sA-D]', text, re.IGNORECASE))

    def _check_author(self, doc: docx.Document) -> bool:
        author = doc.core_properties.author
        return bool(author and len(author) > 2)

    def _check_approval(self, text: str) -> bool:
        approval_keywords = ['approved', 'validated', 'freigegeben', 'release']
        return any(keyword in text.lower() for keyword in approval_keywords)

    def _check_effective_date(self, text: str) -> bool:
        date_patterns = [r'\d{2}\.\d{2}\.\d{4}', r'\d{4}-\d{2}-\d{2}']
        return any(re.search(pattern, text) for pattern in date_patterns)

    def _check_ownership_notice(self, text: str) -> bool:
        ownership_keywords = ['confidential', 'geschäftsgeheimnis', 'business secret', '© ALSTOM']
        return any(keyword.lower() in text.lower() for keyword in ownership_keywords)

    def _check_creating_unit(self, text: str) -> bool:
        alstom_units = ['alstom', 'produktlinie', 'standort', 'dach region']
        return any(unit in text.lower() for unit in alstom_units)

    def _check_confidentiality_level(self, text: str) -> Optional[str]:
        text_lower = text.lower()
        for level in self.confidentiality_levels:
            if level.lower() in text_lower:
                return level
        return None

    def _check_document_numbering(self, text: str) -> bool:
        patterns = list(self.document_numbering_patterns.values())
        return any(re.search(pattern, text) for pattern in patterns)

    def _check_template_compliance(self, doc: docx.Document) -> bool:
        return (doc.core_properties.title is not None and doc.core_properties.author is not None)

    def _check_ams_compliance(self, text: str) -> bool:
        ams_keywords = ['alstom management system', 'ams', 'management handbook']
        return any(keyword in text.lower() for keyword in ams_keywords)


@app.route('/')
def index():
    return render_template('index.html', language='de')

@app.route('/verify', methods=['POST'])
def verify():
    language = request.form.get('language', 'en')

    if 'file' not in request.files:
        return render_template('index.html', error="No file provided", language=language)

    file = request.files['file']

    if file.filename == '':
        return render_template('index.html', error="No file selected", language=language)

    if file.filename.split('.')[-1].lower() not in app.config['ALLOWED_EXTENSIONS']:
        return render_template('index.html', error="Unsupported file format", language=language)

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
    file.save(file_path)

    verifier = AlstomDocumentVerifier()
    results = verifier.verify_document(file_path, language)

    return render_template('index.html', file_name=file.filename, results=results, language=language)


if __name__ == '__main__':
    app.run(debug=True)
